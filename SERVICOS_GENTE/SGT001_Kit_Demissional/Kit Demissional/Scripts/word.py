import docx
import sys
import os
from docx2pdf import convert


def gerarComprovante(titulo, data, agencia, contaCorrente, colaborador, credito):
    try:
        doc = docx.Document("template.docx")

        replaceList = {
            "TITULO": titulo,
            "dd/MM/yyyy": data,
            "AGENCIA": agencia,
            "CONTA-CORRENTE": contaCorrente,
            "COLABORADOR": colaborador,
            "CREDITO": credito
        }

        all_paras = doc.paragraphs
        table = doc.tables[0]

        for para in all_paras:
            inline = para.runs
            for item in inline:
                if 'TITULO' in item.text:
                    item.text = item.text.replace('TITULO', replaceList['TITULO'])

        for row in table.rows:
            for cell in row.cells:
                #print(cell.text, end=' ')
                for element in replaceList:
                    if element in cell.text:
                        cell.text = cell.text.replace(element, replaceList[element])

        doc.save("Comprovante de Pagamento.docx")

        convert("Comprovante de Pagamento.docx")
        convert("Comprovante de Pagamento.docx", "Comprovante de Pagamento.pdf")


        os.remove("Comprovante de Pagamento.docx")

        return "Comprovante gerado com sucesso"
    except Exception as e:
        return e
